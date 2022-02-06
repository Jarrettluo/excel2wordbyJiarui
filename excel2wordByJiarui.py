import time

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import *
from utils import _to_chinese4

XLSX_PATH = "resources/qiu.xlsx"
WORD_PATH = "./tmp/template.docx"

TMP_PATH = "./tmp/template.docx"
OUTPUT_PATH = "./output"

class ReadXlsx:
    def __init__(self, xlsx_path):
        self.xlsx_path = xlsx_path
        self.xlsx_data_list = []

    def read_xlsx(self):
        wb = load_workbook(self.xlsx_path)
        ws = wb.active
        last_proposer = None
        for row in list(ws.rows)[3:]:
            proposal_topic = row[1].value
            proposer = row[2].value
            reporter = row[3].value
            attending_dep = row[4].value
            # 如果提案主题为空，则跳过
            if proposal_topic is None:
                continue
            # 如果其他信息全部为空，则跳过
            if proposer is None and reporter is None:
                continue
            # 判断提案人是否具备，否则沿用之前的信息
            if proposer is None:
                proposer = last_proposer
            else:
                last_proposer = proposer

            row_data_dict = {
                "proposal_topic": proposal_topic,  # 提案主题
                'proposer': proposer,  # 提议人
                'reporter': reporter,  # 报告人
                'attending_dep': attending_dep,  # 列席单位
            }
            self.xlsx_data_list.append(row_data_dict)

        return self.xlsx_data_list


class Write2Word:
    def __init__(self, proposal_list, output_path):
        self.docx_path = TMP_PATH
        self.proposal_list = proposal_list
        self.output_path = output_path

    def open_docx(self):
        doc = Document(self.docx_path)
        # 设置写入文字的字体
        doc.styles['Normal'].font.name = u'方正仿宋简体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋简体')
        doc.styles['Normal'].font.size = Pt(16)
        for i, element in enumerate(self.proposal_list):
            proposal_topic = str(_to_chinese4(i + 1)) + "、" + element['proposal_topic']
            proposer = "提议人：" + element['proposer']
            reporter = "汇报人：" + element['reporter']
            if element['attending_dep'] is None:
                attending_dep = "列席人：" + '无'
            else:
                attending_dep = "列席人：" + element['attending_dep']
            p = doc.add_paragraph()
            p.add_run(proposal_topic).bold = True

            p = doc.add_paragraph()
            p.add_run(proposer).bold = True

            p = doc.add_paragraph()
            p.add_run(reporter).bold = True

            p = doc.add_paragraph()
            p.add_run(attending_dep).bold = True
        # 设置每段缩进2字符
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.first_line_indent = 406400
        doc.save(self.output_path)


def start_switch(input_file, file_name):
    read_xlsx_class = ReadXlsx(input_file)
    data_list = read_xlsx_class.read_xlsx()
    if data_list is None:
        return False
    time_formate = time.strftime('%Y%m%d_%H%M%S', time.localtime())
    output_path = OUTPUT_PATH + "/" + file_name + "_" + time_formate + ".docx"
    word_class = Write2Word(data_list, output_path)
    word_class.open_docx()
    return True


if __name__ == "__main__":
    start_switch("./resources/qiu.xlsx", "qiu_1")
